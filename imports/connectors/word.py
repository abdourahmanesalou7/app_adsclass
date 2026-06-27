"""
Connecteur Word (.docx) : extrait la première table trouvée.
Convention : la 1ère ligne = en-têtes.
"""
from docx import Document
from .base import BaseConnector
from ..security import sanitize_cell, enforce_row_limit, enforce_column_limit


class WordConnector(BaseConnector):
    source_type = 'word'

    def extract(self, filepath, table_index=0, **kwargs):
        doc = Document(filepath)
        if not doc.tables:
            return {'headers': [], 'rows': [], 'meta': {'tables_found': 0}}

        idx = min(table_index, len(doc.tables) - 1)
        table = doc.tables[idx]

        rows_iter = list(table.rows)
        if not rows_iter:
            return {'headers': [], 'rows': []}

        raw_headers = [cell.text for cell in rows_iter[0].cells]
        headers = self._normalize_headers(raw_headers)
        enforce_column_limit(headers)

        data_rows = []
        for row in rows_iter[1:]:
            cells = [sanitize_cell(c.text) for c in row.cells]
            if all((v is None or not str(v).strip()) for v in cells):
                continue
            record = {headers[i]: (cells[i] if i < len(cells) else None) for i in range(len(headers))}
            data_rows.append(record)

        enforce_row_limit(len(data_rows))
        return {
            'headers': headers,
            'rows': data_rows,
            'meta': {'tables_found': len(doc.tables), 'used_table_index': idx},
        }
